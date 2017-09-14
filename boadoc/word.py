import jinja2

import docx
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxtpl import DocxTemplate

from .document import DocFile
from .page import Page, CoverPage, TableOfContent

from .style import PARAGRAPH_KEY

__author__ = 'jbui'


class Writer(DocFile):
    """
    Word document view

    """

    def __init__(self, doc=None, file_path=''):
        """
        Constructor

        :param doc:
        :param file_path:
        :return:
        """
        DocFile.__init__(self)

        if doc:
            # Copy all the content inside document.
            self.__dict__.update(doc.__dict__)

        self.file_path = file_path

        self.jinja_env = jinja2.Environment(
            autoescape=['html', 'xml']
        )

        # Initialize a document to write.
        if doc.template_path:
            # Use a template module.
            self.myDocument = DocxTemplate(doc.template_path)

        else:
            self.myDocument = docx.Document()

    def write(self):
        """
        Write the driver.

        :return:
        """
        try:
            if self.context:
                self.myDocument.render(self.context, self.jinja_env)

            for page in self.pages:
                page.write(self)
        except Exception as e:
            import wx

            text = "Problem writing the documents.  \n\n" + str(e)

            dlg = wx.MessageBox(message=text, caption="Error")

        finally:
            self.close()

    def close(self):
        """
        Close and save the document.

        :return:
        """
        self.myDocument.save(self.file_path)

    def write_paragraph(self, texts, **kwargs):
        """
        Return a paragraph newly added to the end of the document, populated with text and having paragraph style
        style. text can contain tab (\t) characters, which are converted to the appropriate XML form for a tab. text
        can also include newline (\n) or carriage return (\r) characters, each of which is converted to a line break.

        :param texts:
        :param kwargs:
        :return:
        """
        style = kwargs.get('style')

        if PARAGRAPH_KEY.get(style):
            return self.myDocument.add_paragraph(texts, style=PARAGRAPH_KEY[style][0])
        else:
            return self.myDocument.add_paragraph(texts)

    def write_heading(self, text, level=1):
        """
        Return a heading paragraph newly added to the end of the document, populated with text and having the heading
        paragraph style determined by level. If level is 0, the style is set to 'Title'. If level is 1
        (or not present), 'Heading1' is used. Otherwise the style is set to 'Heading{level}'. If level is outside the
        range 0-9, ValueError is raised.

        :param text:
        :param level: level of the heading inside the document.
        """
        return self.myDocument.add_heading(text, level=level)

    def write_table(self, data, **kwargs):
        """
        Write table

        :param data:
        :param kwargs
        :return:
        """
        rows = len(data)
        cols = len(data[0])

        table = self.myDocument.add_table(rows=rows, cols=cols)

        for i_row in range(0, rows):
            cells = table.rows[i_row].cells
            for i_cell in range(0, cols):
                cells[i_cell].text = str(data[i_row][i_cell])

    def write_section(self, start_type=2):
        """
        Return a Section object representing a new section added at the end of the document. The optional start_type
        argument must be a member of the WD_SECTION_START enumeration defaulting to WD_SECTION.NEW_PAGE if not provided.

        :param start_type:
        :return:
        """
        return self.myDocument.add_section()

    def write_page(self, page):
        """
        Add page break.

        :param page:
        :return:
        """
        pass

    def write_picture(self, image_path_or_stream='', **kwargs):
        """
        Return a new picture shape added in its own paragraph at the end of the document. The picture contains the
        image at image_path_or_stream, scaled based on width and height. If neither width nor height is specified,
        the picture appears at its native size. If only one is specified, it is used to compute a scaling factor that
        is then applied to the unspecified dimension, preserving the aspect ratio of the image. The native size of the
        picture is calculated using the dots-per-inch (dpi) value specified in the image file, defaulting to 72 dpi if
        no value is specified, as is often the case.

        :param image_path_or_stream:
        :param **kwargs:
            width = float value in inches
            height = float value in inches
        :return:
        """
        width = kwargs.get('width')
        height = kwargs.get('height')

        image = None
        if width and height:
            image = self.myDocument.add_picture(image_path_or_stream, width=Inches(width), height=Inches(height))
        elif width:
            image = self.myDocument.add_picture(image_path_or_stream, width=Inches(width))
        elif height:
            image = self.myDocument.add_picture(image_path_or_stream, height=Inches(height))
        else:
            image = self.myDocument.add_picture(image_path_or_stream)

        return image

    def add_table(self, **kwargs):
        """
        Add a table having row and column counts of rows and cols respectively and table style of style.
        If style is None, a table with no style is produced.

        :param **kwargs:
            rows: Number of rows.
            cols: Number of cols.
            style: Table style.
            alignment: Word allows a table to be aligned between the page margins either left, right, or center.
            allow_autofit: Word has two algorithms for laying out a table, fixed-width or autofit.
        :return:
        """
        rows = kwargs.get('rows')
        cols = kwargs.get('cols')

        table = None

        if rows and cols:
            table = self.myDocument.add_table(rows, cols)

        # If there is a table.
        if table:
            if kwargs.get('style'):
                table.style = kwargs.get('style')

            if kwargs.get('alignment'):
                if kwargs.get('alignment') == self.TABLE_ALIGNMENT_LEFT:
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                elif kwargs.get('alignment') == self.TABLE_ALIGNMENT_CENTER:
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif kwargs.get('alignment') == self.TABLE_ALIGNMENT_RIGHT:
                    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

            if kwargs.get('allow_autofit'):
                table.allow_autofit = kwargs.get('allow_autofit')

        return table

    def add_new_row(self, table, row):
        """

        :param table:
        :param row:
        :return:
        """
        row_cell = table.add_row().cells

        for i in range(0, len(row)):
            row_cell[i].text = str(row[i])

        return True

    def add_row_in_index(self, table, row, row_index):
        """

        :param table:
        :param row:
        :param row_index:
        :return:
        """
        cells = table.rows[row_index].cells

        for i in range(0, len(row)):
            cells[i].text = str(row[i])

        return True

    def write_page_break(self):
        """
        
        :return: 
        """
        return self.myDocument.add_page_break()

    def open_document(self, filename):
        """
        Open a old document.

        :param filename:
        """
        self.myDocument = docx.Document(filename)

    def save(self, file_name):
        """
        Save the document to file name.

        :param file_name:
        """
        self.myDocument.save(file_name)
